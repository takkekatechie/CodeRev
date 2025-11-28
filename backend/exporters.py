"""
Exporters for CodeReviewPro scan results
"""
import json
import base64
import io
import pandas as pd
from typing import Dict, Any, List, Union

# Import optional dependencies
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
except ImportError:
    SimpleDocTemplate = None

def _flatten_issues(results: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Flatten issues for tabular formats"""
    issues = results.get('issues', [])
    flat_issues = []
    
    for issue in issues:
        flat_issues.append({
            'Category': issue.get('category', 'Unknown'),
            'Severity': issue.get('severity', 'Unknown'),
            'File': issue.get('filePath', ''),
            'Line Start': issue.get('lineStart', ''),
            'Line End': issue.get('lineEnd', ''),
            'Description': issue.get('description', ''),
            'Recommendation': issue.get('recommendation', '')
        })
    
    return flat_issues

def export_to_csv(results: Dict[str, Any]) -> str:
    """Export results to CSV string"""
    flat_issues = _flatten_issues(results)
    if not flat_issues:
        return "No issues found"
    
    df = pd.DataFrame(flat_issues)
    return df.to_csv(index=False)

def export_to_excel(results: Dict[str, Any]) -> bytes:
    """Export results to Excel bytes"""
    flat_issues = _flatten_issues(results)
    output = io.BytesIO()
    
    if not flat_issues:
        # Create empty dataframe with columns
        df = pd.DataFrame(columns=['Category', 'Severity', 'File', 'Line Start', 'Line End', 'Description', 'Recommendation'])
    else:
        df = pd.DataFrame(flat_issues)
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Issues')
        
        # Auto-adjust column widths (basic approximation)
        worksheet = writer.sheets['Issues']
        for idx, col in enumerate(df.columns):
            max_len = max(
                df[col].astype(str).map(len).max(),
                len(col)
            ) + 2
            # Cap width at 50
            worksheet.column_dimensions[chr(65 + idx)].width = min(max_len, 50)
            
    return output.getvalue()

def export_to_text(results: Dict[str, Any]) -> str:
    """Export results to plain text"""
    issues = results.get('issues', [])
    summary = results.get('summary', {})
    
    lines = []
    lines.append("CodeReviewPro Scan Report")
    lines.append("=========================")
    lines.append(f"Total Issues: {summary.get('totalIssues', 0)}")
    lines.append(f"Scan Date: {results.get('scanDate', 'Unknown')}")
    lines.append("")
    
    if not issues:
        lines.append("No issues found.")
        return "\n".join(lines)
    
    # Group by category
    by_category = {}
    for issue in issues:
        cat = issue.get('category', 'Other')
        if cat not in by_category:
            by_category[cat] = []
        by_category[cat].append(issue)
        
    for cat, cat_issues in by_category.items():
        lines.append(f"Category: {cat.upper()}")
        lines.append("-" * (len(cat) + 10))
        
        for issue in cat_issues:
            lines.append(f"[{issue.get('severity', 'INFO').upper()}] {issue.get('filePath', '')}:{issue.get('lineStart', '?')}")
            lines.append(f"Description: {issue.get('description', '')}")
            if issue.get('recommendation'):
                lines.append(f"Recommendation: {issue.get('recommendation', '')}")
            lines.append("")
        lines.append("")
        
    return "\n".join(lines)

def export_to_word(results: Dict[str, Any]) -> bytes:
    """Export results to Word document bytes"""
    if Document is None:
        raise ImportError("python-docx is not installed")
        
    doc = Document()
    doc.add_heading('CodeReviewPro Scan Report', 0)
    
    summary = results.get('summary', {})
    doc.add_paragraph(f"Total Issues: {summary.get('totalIssues', 0)}")
    doc.add_paragraph(f"Scan Date: {results.get('scanDate', 'Unknown')}")
    
    issues = results.get('issues', [])
    if not issues:
        doc.add_paragraph("No issues found.")
    else:
        # Group by category
        by_category = {}
        for issue in issues:
            cat = issue.get('category', 'Other')
            if cat not in by_category:
                by_category[cat] = []
            by_category[cat].append(issue)
            
        for cat, cat_issues in by_category.items():
            doc.add_heading(cat.capitalize(), level=1)
            
            for issue in cat_issues:
                p = doc.add_paragraph()
                runner = p.add_run(f"[{issue.get('severity', 'INFO').upper()}] ")
                if issue.get('severity', '').lower() == 'error':
                    runner.font.color.rgb = RGBColor(255, 0, 0)
                elif issue.get('severity', '').lower() == 'warning':
                    runner.font.color.rgb = RGBColor(255, 165, 0)
                
                p.add_run(f"{issue.get('filePath', '')} (Line {issue.get('lineStart', '?')})").bold = True
                
                doc.add_paragraph(issue.get('description', ''))
                
                if issue.get('recommendation'):
                    rec_p = doc.add_paragraph()
                    rec_p.add_run("Recommendation: ").bold = True
                    rec_p.add_run(issue.get('recommendation', ''))
                
                doc.add_paragraph() # Spacer

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def export_to_pdf(results: Dict[str, Any]) -> bytes:
    """Export results to PDF bytes"""
    if SimpleDocTemplate is None:
        raise ImportError("reportlab is not installed")
        
    output = io.BytesIO()
    doc = SimpleDocTemplate(output, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    story.append(Paragraph("CodeReviewPro Scan Report", styles['Title']))
    story.append(Spacer(1, 12))
    
    # Summary
    summary = results.get('summary', {})
    story.append(Paragraph(f"Total Issues: {summary.get('totalIssues', 0)}", styles['Normal']))
    story.append(Paragraph(f"Scan Date: {results.get('scanDate', 'Unknown')}", styles['Normal']))
    story.append(Spacer(1, 12))
    
    issues = results.get('issues', [])
    if not issues:
        story.append(Paragraph("No issues found.", styles['Normal']))
    else:
        # Group by category
        by_category = {}
        for issue in issues:
            cat = issue.get('category', 'Other')
            if cat not in by_category:
                by_category[cat] = []
            by_category[cat].append(issue)
            
        for cat, cat_issues in by_category.items():
            story.append(Paragraph(cat.capitalize(), styles['Heading2']))
            
            for issue in cat_issues:
                sev = issue.get('severity', 'INFO').upper()
                color = "black"
                if sev == 'ERROR': color = "red"
                elif sev == 'WARNING': color = "orange"
                elif sev == 'INFO': color = "blue"
                
                header_text = f"<font color='{color}'><b>[{sev}]</b></font> <b>{issue.get('filePath', '')}</b> (Line {issue.get('lineStart', '?')})"
                story.append(Paragraph(header_text, styles['Normal']))
                
                story.append(Paragraph(issue.get('description', ''), styles['Normal']))
                
                if issue.get('recommendation'):
                    rec_text = f"<b>Recommendation:</b> {issue.get('recommendation', '')}"
                    story.append(Paragraph(rec_text, styles['Normal']))
                
                story.append(Spacer(1, 6))
    
    doc.build(story)
    return output.getvalue()

def export_report(results: Dict[str, Any], format_type: str) -> Dict[str, Any]:
    """
    Export report to specified format.
    Returns dict with 'content' (str or base64) and 'is_binary' (bool)
    """
    format_type = format_type.lower()
    
    if format_type == 'csv':
        return {'content': export_to_csv(results), 'is_binary': False}
        
    elif format_type == 'excel':
        content = export_to_excel(results)
        return {'content': base64.b64encode(content).decode('utf-8'), 'is_binary': True}
        
    elif format_type == 'text':
        return {'content': export_to_text(results), 'is_binary': False}
        
    elif format_type == 'word':
        content = export_to_word(results)
        return {'content': base64.b64encode(content).decode('utf-8'), 'is_binary': True}
        
    elif format_type == 'pdf':
        content = export_to_pdf(results)
        return {'content': base64.b64encode(content).decode('utf-8'), 'is_binary': True}
        
    elif format_type == 'markdown':
        # We can reuse text format or make a specific one, for now text is close enough or we can implement markdown specific
        # Reusing text but maybe wrapping in code blocks? No, let's just use text for now or implement simple markdown
        return {'content': export_to_text(results), 'is_binary': False}
        
    else:
        raise ValueError(f"Unsupported format: {format_type}")
